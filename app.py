import PySimpleGUI as sg
from docx import Document
import pandas as pd
import os


sg.theme('Reddit')    # Keep things interesting for your users
layout = [
    [sg.Text('Selecione o arquivo XLSX:'),
     sg.InputText(key='-FILE-'), sg.FileBrowse()],
    [sg.Text('Selecione a pasta onde será armazenado:'),
     sg.InputText(key='-FOLDER-'), sg.FolderBrowse()],
    [sg.Button('Extrair'), sg.Button('Cancelar')]
]

window = sg.Window('Extrair informações do arquivo XLSX', layout)

while True:
    event, values = window.read()

    if event == sg.WIN_CLOSED or event == 'Cancelar':
        break
    elif event == 'Extrair':
        if values['-FILE-'] and values['-FOLDER-']:
            file_path = values['-FILE-']
            output_dir = values['-FOLDER-']

            try:
                df = pd.read_excel(file_path)
                print(df)

                for index, row in df.iterrows():
                    file_name = f"{output_dir}/{row['nome']}.docx"
                    doc = Document()
                    doc.add_heading('Informações', level=1)
                    doc.add_paragraph(f"Nome: {row['nome']}")
                    doc.add_paragraph(f"Estado Civil: {row['estado civil']}")
                    doc.add_paragraph(f"Profissão: {row['profissão']}")
                    doc.add_paragraph(f"Valores: {row['valores']}")
                    doc.save(file_name)

                sg.popup('Arquivo carregado com sucesso!',
                         f'O arquivo tem {len(df)} linhas.')
            except Exception as e:
                sg.popup_error(f'Ocorreu um erro ao carregar o arquivo: {e}')
        else:
            sg.popup_error('Por favor, selecione um arquivo XLSX.')

window.close()
